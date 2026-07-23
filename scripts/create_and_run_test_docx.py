#!/usr/bin/env python3
"""
Create a small test DOCX with headings and paragraphs and run the docx_to_yaml converter.
"""
import os, sys
from docx import Document
from pathlib import Path

here = Path.cwd()
run_dir = here / 'run'
run_dir.mkdir(exist_ok=True)

docx_path = run_dir / 'test_docx_for_conversion.docx'
yaml_out = run_dir / 'test_docx_for_conversion.yaml'

# create docx
doc = Document()
doc.add_heading('Test Document Title', level=1)
doc.add_paragraph('This is the first paragraph of the test document.')
doc.add_heading('Section One', level=2)
doc.add_paragraph('Some more text in section one.')
doc.add_heading('Subsection', level=3)
doc.add_paragraph('Details in subsection.')

doc.save(docx_path)
print(f'Wrote test DOCX to {docx_path}')

# run the converter
py = sys.executable
converter = here / 'scripts' / 'docx_to_yaml.py'
if not converter.exists():
    print('Converter script not found:', converter)
    sys.exit(2)

cmd = [py, str(converter), str(docx_path), str(yaml_out)]
print('Running:', ' '.join(cmd))
import subprocess
res = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
print('Exit', res.returncode)
print(res.stdout)
if res.stderr:
    print('Stderr:', res.stderr)

if yaml_out.exists():
    print('YAML created at', yaml_out)
else:
    print('YAML not created')

print('Done')
